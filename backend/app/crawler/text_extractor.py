"""
Utilities to extract clean plain text from HTML, PDF, and DOCX content.
Hebrew text is preserved as-is (Unicode). Tags / boilerplate are stripped.
"""

import re
import io
import logging
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)


# ─── HTML ────────────────────────────────────────────────────────────────────

def extract_text_from_html(html: str, base_url: str = "") -> tuple[str, str]:
    """Parse HTML and return (title, body_text)."""
    soup = BeautifulSoup(html, "lxml")

    for tag in soup.find_all(
        ["script", "style", "nav", "header", "footer", "noscript", "iframe", "svg"]
    ):
        tag.decompose()

    title = ""
    og_title = soup.find("meta", property="og:title")
    if og_title and og_title.get("content"):
        title = og_title["content"].strip()
    elif soup.title and soup.title.string:
        title = soup.title.string.strip()
    else:
        h1 = soup.find("h1")
        if h1:
            title = h1.get_text(strip=True)

    main = (
        soup.find("main")
        or soup.find("article")
        or soup.find(id="content")
        or soup.find(class_=re.compile(r"content|main|article", re.I))
        or soup.body
    )

    body_text = ""
    if main:
        body_text = main.get_text(separator=" ", strip=True)

    body_text = re.sub(r"\s{2,}", " ", body_text).strip()
    return title, body_text


# ─── PDF ─────────────────────────────────────────────────────────────────────

def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extract plain text from a PDF binary using pypdf."""
    try:
        import pypdf

        reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
        pages_text = []
        for page in reader.pages:
            text = page.extract_text() or ""
            pages_text.append(text)
        full_text = "\n".join(pages_text)
        return re.sub(r"\s{2,}", " ", full_text).strip()
    except ImportError:
        logger.warning("pypdf not installed — PDF text extraction skipped")
        return ""
    except Exception as e:
        logger.warning(f"PDF extraction error: {e}")
        return ""


# ─── DOCX ────────────────────────────────────────────────────────────────────

def extract_text_from_docx(docx_bytes: bytes) -> str:
    """Extract plain text from a DOCX binary using python-docx."""
    try:
        import docx as python_docx  # python-docx package

        doc = python_docx.Document(io.BytesIO(docx_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        full_text = "\n".join(paragraphs)
        return re.sub(r"\s{2,}", " ", full_text).strip()
    except ImportError:
        logger.warning("python-docx not installed — DOCX text extraction skipped")
        return ""
    except Exception as e:
        logger.warning(f"DOCX extraction error: {e}")
        return ""


# ─── Generic ─────────────────────────────────────────────────────────────────

def extract_text_from_bytes(content: bytes, url: str) -> str:
    """Detect file type from URL and extract text accordingly."""
    url_lower = url.lower()
    if url_lower.endswith(".pdf"):
        return extract_text_from_pdf(content)
    if url_lower.endswith(".docx"):
        return extract_text_from_docx(content)
    # Fallback: treat as HTML
    try:
        html = content.decode("utf-8", errors="replace")
        _, text = extract_text_from_html(html)
        return text
    except Exception:
        return ""


def clean_text(text: str) -> str:
    """Normalize text: collapse whitespace, strip control characters."""
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    text = re.sub(r" {2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
